from docx import Document
import os
import pandas as pd
import re

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement, ns

def normalizar(texto):
    if texto is None:
        return ""

    return (
        str(texto)
        .replace("\xa0", " ")
        .replace("\n", " ")
        .strip()
    )

COLUMNAS_EID = {

    # =========================
    # BLOQUE 1
    # =========================
    1: {
        "eid": "Estándares Indicativos de Desempeño asociado",
        "capacidad": "Capacidad abordada en la sesión de asesoría",
        "dimension": "Dimensión asociada al EID seleccionado",
        "practica": "Indique qué práctica se está abordando en el establecimiento a partir del EID trabajado.",
        "segunda": "¿Se trabajó una segunda capacidad o estándar en su sesión de asesoría?",
    },

    # =========================
    # BLOQUE 2
    # =========================
    2: {
        "capacidad": "Capacidad (2) abordada en la sesión de asesoría",
        "dimension": "Dimensión asociada al EID seleccionado2",
        "practica": "Indique qué práctica (2) se está abordando en el establecimiento a partir del EID trabajado.",
    },

}

def obtener_bloque_eid2(fila, columnas_df):

    segunda = valor_visible(
        fila.get(
            COLUMNAS_EID[1]["segunda"]
        )
    )

    if normalizar(segunda).upper() != "SÍ":
        return None

    capacidad = valor_visible(
        fila.get(
            COLUMNAS_EID[2]["capacidad"]
        )
    )

    dimension = valor_visible(
        fila.get(
            COLUMNAS_EID[2]["dimension"]
        )
    )

    practica = valor_visible(
        fila.get(
            COLUMNAS_EID[2]["practica"]
        )
    )

    subdimension, estandar = obtener_subdimension_y_estandar(
        fila,
        columnas_df,
        bloque=2
    )

    return {
        "capacidad": capacidad,
        "dimension": dimension,
        "subdimension": subdimension,
        "estandar": estandar,
        "practica": practica,
        "segunda": segunda,
    }
    
def buscar_valor_en_bloques(fila, columnas_df, texto_base):
    """
    Busca un valor en todas las variantes del campo (1,2,3,4...)
    """

    for col in columnas_df:
        col_lower = str(col).lower()

        if texto_base.lower() in col_lower:
            valor = valor_visible(fila.get(col))

            if valor:
                return valor

    return ""

def obtener_subdimension_y_estandar(fila, columnas_df, bloque=1):

    subdimension = ""
    estandar = ""

    # =========================
    # DIMENSIÓN
    # =========================
    col_dimension = COLUMNAS_EID[bloque]["dimension"]

    dimension = valor_visible(fila.get(col_dimension))

    if not dimension:
        return "", ""

    dimension_lower = normalizar(dimension).lower()

    # =========================
    # SUBDIMENSIÓN
    # =========================
    for col in columnas_df:

        col_lower = normalizar(col).lower()

        if (
            "sub dimensión" in col_lower
            and dimension_lower in col_lower
        ):

            # bloque 1 = sin número
            if bloque == 1 and not col_lower.endswith("2"):
                valor = valor_visible(fila.get(col))

                if valor:
                    subdimension = valor
                    break

            # bloque 2 = columnas terminadas en 2
            elif bloque == 2 and col_lower.endswith("2"):

                valor = valor_visible(fila.get(col))

                if valor:
                    subdimension = valor
                    break

    # =========================
    # ESTÁNDAR
    # =========================
    if subdimension:

        sub_lower = normalizar(subdimension).lower()

        for col in columnas_df:

            col_lower = normalizar(col).lower()

            if (
                "estándar asociado" in col_lower
                and sub_lower in col_lower
            ):

                # bloque 1
                if bloque == 1 and not col_lower.endswith("2"):

                    valor = valor_visible(fila.get(col))

                    if valor:
                        estandar = valor
                        break

                # bloque 2
                elif bloque == 2 and col_lower.endswith("2"):

                    valor = valor_visible(fila.get(col))

                    if valor:
                        estandar = valor
                        break

    return subdimension, estandar

def buscar_columna(columna_objetivo, columnas_df):
    """
    Busca una columna real en el DataFrame usando coincidencia parcial.
    """
    objetivo = columna_objetivo.lower().replace("\xa0", " ").strip()

    for col in columnas_df:
        col_norm = str(col).lower().replace("\xa0", " ").strip()

        if objetivo in col_norm:
            return col

    return None

def aplicar_color_fondo(celda, color_hex="D9E1F2"):
    """
    Aplica color de fondo a una celda de tabla Word.
    """
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(ns.qn('w:val'), 'clear')
    shd.set(ns.qn('w:color'), 'auto')
    shd.set(ns.qn('w:fill'), color_hex)
    tc_pr.append(shd)


def agregar_titulo_y_logo(doc, ruta_logo="logo_mineduc.png"):
    """
    Agrega el logo institucional y el título del informe
    al inicio del documento.
    """

    # Logo
    if os.path.exists(ruta_logo):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(ruta_logo, width=Cm(4))

    # Espacio
    doc.add_paragraph("")

    # Título
    titulo = doc.add_heading(
        "Informe Individual de Registro de Asesoría MINEDUC",
        level=0
    )
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Separación visual
    doc.add_paragraph("")

def limpiar_documento(doc):
    """
    Limpia el contenido del documento Word (párrafos y tablas),
    pero conserva la sección para evitar errores de python-docx.
    """
    body = doc._element.body
    for element in list(body):
        # Mantener la definición de sección (<w:sectPr>)
        if element.tag.endswith('sectPr'):
            continue
        body.remove(element)

# =====================================================
# UTILIDADES
# =====================================================
def normalizar_texto(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def valor_visible(valor):
    if pd.isna(valor):
        return ""

    # Convertir a string SIN strip agresivo
    texto = str(valor)

    # Reemplazar caracteres invisibles comunes
    texto = texto.replace('\r\n', '\n').replace('\r', '\n')

    # Si sigue siendo sólo espacios o saltos, lo tratamos como vacío
    if texto.strip() == "":
        return ""

    return texto

def limpiar_nombre_archivo(nombre):
    return re.sub(r'[\\/*?:"<>|]', "", str(nombre)).strip()


# =====================================================
# PLANIFICACIÓN
# =====================================================
def obtener_objetivos_planificacion(
    df_planificacion, region, deprov, modalidad, nombre_asesoria
):

    filtro = (
        (df_planificacion["Indique su región"].str.upper().str.strip() == region.upper().strip()) &
        (df_planificacion["Deprov"].str.upper().str.strip() == deprov.upper().strip()) &
        (df_planificacion["Tipo Asesoría"].str.upper().str.strip() == modalidad.upper().strip()) &
        (df_planificacion["Nombre Asesoría"].str.upper().str.strip() == nombre_asesoria.upper().strip())
    )

    fila = df_planificacion[filtro]

    if fila.empty:
        return "No informado", "No informado"

    return (
        normalizar_texto(fila.iloc[0].get("Objetivo estratégico de la asesoría")),
        normalizar_texto(fila.iloc[0].get("Objetivo anual de la asesoría")),
    )


# =====================================================
# SECCIONES DEL INFORME
# =====================================================
def agregar_encabezado(doc, region, deprov, modalidad, nombre_asesoria,
                       objetivo_estrategico, objetivo_anual):

    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = "Table Grid"

    filas = [
        ("Región", region),
        ("DEPROV", deprov),
        ("Modalidad", modalidad),
        ("RBD / Nombre de la Asesoría", nombre_asesoria),
        ("Objetivo estratégico de la asesoría", objetivo_estrategico),
        ("Objetivo anual de la asesoría", objetivo_anual),
    ]

    for i, (k, v) in enumerate(filas):
        tabla.cell(i, 0).text = k
        tabla.cell(i, 1).text = v


def agregar_antecedentes_generales(doc, datos):

    doc.add_heading("ANTECEDENTES GENERALES", level=1)

    tabla = doc.add_table(rows=1, cols=8)
    tabla.style = "Table Grid"

    encabezados = [
        "N° de sesión", "Supervisor", "Fecha de la reunión",
        "Hora inicio", "Hora término", "Tipo de encuentro",
        "Participantes", "Objetivo de la sesión"
    ]

    
    for i, h in enumerate(encabezados):
        celda = tabla.rows[0].cells[i]
        celda.text = h
        aplicar_color_fondo(celda)

    datos = datos.sort_values("NUM SESIÓN")

    for _, fila in datos.iterrows():
        r = tabla.add_row().cells
        r[0].text = str(fila.get("NUM SESIÓN", ""))
        r[1].text = valor_visible(fila.get("Supervisor"))
        r[2].text = valor_visible(fila.get("Fecha de la reunión"))       
        r[3].text = valor_visible(
            fila.get("Indique hora aproximada de inicio de la reunión")
            or fila.get("Hora de inicio")
        )
        r[4].text = valor_visible(
            fila.get("Indique hora aproximada de término de la reunión")
            or fila.get("Hora de finalización")
        )
        r[5].text = valor_visible(fila.get("Tipo de encuentro"))
        r[6].text = valor_visible(fila.get("Participantes de la reunión"))
        r[7].text = valor_visible(fila.get("Objetivo de la sesión de asesoría"))


def agregar_eid_capacidades_practicas(doc, datos):

    columnas = [
        ("N° de sesión", "NUM SESIÓN"),
        ("Estándares Indicativos de Desempeño asociado",
         "Estándares Indicativos de Desempeño asociado"),
        ("Capacidad abordada en la sesión de asesoría",
         "Capacidad abordada en la sesión de asesoría"),
        ("Dimensión asociada al EID seleccionado",
         "Dimensión asociada al EID seleccionado"),
        ("Sub dimensión", None),
        ("Estándar asociado a la sub dimensión", None),
        ("Práctica abordada",
         "Indique qué práctica se está abordando en el establecimiento a partir del EID trabajado."),
        ("¿Se trabajó una segunda capacidad o estándar?",
         "¿Se trabajó una segunda capacidad o estándar en su sesión de asesoría?"),
    ]

    # ✅ EID 1
    doc.add_heading("EID, CAPACIDADES Y PRÁCTICAS ABORDADAS (1)", level=1)

    tabla = doc.add_table(rows=1, cols=len(columnas))
    tabla.style = "Table Grid"

    for i, (titulo, _) in enumerate(columnas):
        c = tabla.rows[0].cells[i]
        c.text = titulo
        aplicar_color_fondo(c)

    # ✅ recorrer filas UNA sola vez
    for _, fila in datos.iterrows():

        r = tabla.add_row().cells

        subdimension, estandar = obtener_subdimension_y_estandar(
            fila, datos.columns
        )

        for i, (titulo, col) in enumerate(columnas):

            if titulo == "Sub dimensión":
                r[i].text = valor_visible(subdimension)

            elif titulo == "Estándar asociado a la sub dimensión":
                r[i].text = valor_visible(estandar)

            elif col == "Capacidad abordada en la sesión de asesoría":
                r[i].text = valor_visible(
                    buscar_valor_en_bloques(
                        fila, datos.columns,
                        "Capacidad abordada en la sesión de asesoría"
                    )
                )

            elif col == "Dimensión asociada al EID seleccionado":
                r[i].text = valor_visible(
                    buscar_valor_en_bloques(
                        fila, datos.columns,
                        "Dimensión asociada al EID seleccionado"
                    )
                )

            elif "práctica" in str(col).lower():
                r[i].text = valor_visible(
                    buscar_valor_en_bloques(
                        fila, datos.columns,
                        "práctica se está abordando"
                    )
                )

            elif "segunda capacidad" in str(col).lower():
                r[i].text = valor_visible(
                    buscar_valor_en_bloques(
                        fila, datos.columns,
                        "segunda capacidad o estándar"
                    )
                )

            else:
                r[i].text = valor_visible(fila.get(col))

        # ✅ AHORA -> EID 2 EN MISMO LOOP (CORRECTO)
        bloque2 = obtener_bloque_eid2(fila, datos.columns)

        if bloque2:   # ✅ si existe, lo mostramos

            doc.add_heading("EID, CAPACIDADES Y PRÁCTICAS ABORDADAS (2)", level=1)

            tabla2 = doc.add_table(rows=1, cols=7)
            tabla2.style = "Table Grid"

            headers2 = [
                "N° de sesión",
                "Capacidad abordada",
                "Dimensión",
                "Sub dimensión",
                "Estándar",
                "Práctica",
                "¿Segunda capacidad?"
            ]

            for j, h in enumerate(headers2):
                c = tabla2.rows[0].cells[j]
                c.text = h
                aplicar_color_fondo(c)

            r2 = tabla2.add_row().cells

            r2[0].text = valor_visible(fila.get("NUM SESIÓN"))
            r2[1].text = bloque2["capacidad"]
            r2[2].text = bloque2["dimension"]
            r2[3].text = bloque2["subdimension"]
            r2[4].text = bloque2["estandar"]
            r2[5].text = bloque2["practica"]
            r2[6].text = bloque2["segunda"]

def agregar_otras_indicaciones(doc, datos):

    columnas = [
        (
            "Estrategia / acciones de acompañamiento realizadas",
            ["estrategia"],
        ),
        (
            "Tema no planificado y cómo fue abordado",
            ["tema no planificado"],
        ),
        (
            "Cambios o progresos evidenciados",
            ["cambios o progresos"],
        ),
        (
            "Dificultades identificadas",
            ["dificultades"],
        ),
        (
            "Acuerdos concretos de la reunión",
            ["acuerdos concretos"],
        ),
        (
            "Próximos pasos y responsables",
            ["próximos pasos"],
        ),
        (
            "Comentarios u observaciones",
            ["comentarios"],
        ),
    ]

    # Verificar si existe información
    existe_info = False
    for _, posibles in columnas:
        for col_ref in posibles:
            if buscar_columna(col_ref, datos.columns):
                existe_info = True
                break
        if existe_info:
            break

    if not existe_info:
        return

    doc.add_heading("OTRAS INDICACIONES", level=1)

    tabla = doc.add_table(rows=1, cols=len(columnas) + 1)
    tabla.style = "Table Grid"

    # Encabezado
    tabla.rows[0].cells[0].text = "N° sesión"
    aplicar_color_fondo(tabla.rows[0].cells[0])

    for i, (titulo, _) in enumerate(columnas):
        celda = tabla.rows[0].cells[i + 1]
        celda.text = titulo
        aplicar_color_fondo(celda)

    # Filas
    for _, fila in datos.iterrows():
        r = tabla.add_row().cells
        r[0].text = str(fila.get("NUM SESIÓN", ""))

        for i, (_, posibles) in enumerate(columnas):

            valor = ""  # ✅ siempre inicializado

            for col_ref in posibles:
                col_real = buscar_columna(col_ref, datos.columns)

                if col_real:
                    texto = valor_visible(fila.get(col_real))
                    if texto:
                        valor = texto
                        break

            r[i + 1].text = valor

# =====================================================
# FUNCIÓN PRINCIPAL
# =====================================================
def generar_informes_registro(df_registros, df_planificacion,
                              carpeta_salida, plantilla="plantilla_registro.docx"):

        # Normalizar nombres de columnas del registro
    df_registros.columns = (
        df_registros.columns
        .astype(str)
        .str.strip()
        .str.replace('\u00a0', ' ')   # espacios duros
    )
    
    df_planificacion.columns = (
        df_planificacion.columns
        .astype(str)
        .str.strip()
        .str.replace('\u00a0', ' ')
    )

    os.makedirs(carpeta_salida, exist_ok=True)

    modalidad_mapa = {
        "Directa EE": "Directa a Establecimientos",
        "Directa a Sostenedor": "Directa a Sostenedor",
        "Red EE": "Red de Establecimientos",
        "Red de Sostenedor": "Red de Sostenedor"
    }

    grupos = df_registros.groupby(
        ["INDIQUE SU REGIÓN", "Deprov", "Modalidad Asesoría", "Nombre Asesoría"]
    )

    for (region, deprov, modalidad, nombre), datos in grupos:

        region = normalizar_texto(region)
        deprov = normalizar_texto(deprov)
        modalidad = normalizar_texto(modalidad)
        nombre = normalizar_texto(nombre)

        obj_est, obj_anual = obtener_objetivos_planificacion(
            df_planificacion, region, deprov, modalidad, nombre
        )

        doc = Document(plantilla)
        limpiar_documento(doc)
        
        # Título institucional y logo
        agregar_titulo_y_logo(doc)
        
        # Contenido del informe
        agregar_encabezado(doc, region, deprov, modalidad, nombre, obj_est, obj_anual)

        agregar_antecedentes_generales(doc, datos)
        agregar_eid_capacidades_practicas(doc, datos)
        agregar_otras_indicaciones(doc, datos)

        carpeta_region = limpiar_nombre_archivo(region)
        carpeta_deprov = limpiar_nombre_archivo(deprov)
        carpeta_modalidad = modalidad_mapa.get(modalidad, limpiar_nombre_archivo(modalidad))

        ruta_final = os.path.join(
            carpeta_salida,
            carpeta_region,
            carpeta_deprov,
            carpeta_modalidad
        )

        os.makedirs(ruta_final, exist_ok=True)

        archivo = limpiar_nombre_archivo(
            f"Informe_Registro_{region}_{deprov}_{modalidad}_{nombre}.docx"
        )

        doc.save(os.path.join(ruta_final, archivo))
